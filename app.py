import streamlit as st
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from concurrent.futures import ThreadPoolExecutor, as_completed
import pandas as pd
from io import BytesIO
from docx import Document
import re
from collections import defaultdict
import time

st.set_page_config(page_title="Enterprise IA & Tech Audit Tool", layout="wide")

# -----------------------------
# 1. URL NORMALIZER & FILTER
# -----------------------------
def normalize_url(url):
    parsed = urlparse(url)
    clean = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
    if clean != parsed.scheme + "://" + parsed.netloc + "/":
        clean = clean.rstrip('/')
    return clean

# -----------------------------
# 2. MEMORY-OPTIMIZED THREAD WORKER
# -----------------------------
def fetch_and_analyze(url, domain, integration_patterns):
    links = set()
    integrations_found = set()
    forms_found = []
    has_calculator = False
    
    # FILTER: Ignore non-HTML files (PDFs, images) to save massive amounts of time and RAM
    if any(url.lower().endswith(ext) for ext in ['.pdf', '.jpg', '.jpeg', '.png', '.gif', '.zip', '.xml', '.css', '.js', '.doc', '.docx']):
        return url, set(), set(), [], False

    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        res = requests.get(url, timeout=10, headers=headers)
        
        # Only parse if it's actually an HTML page
        if 'text/html' not in res.headers.get('Content-Type', ''):
            return url, set(), set(), [], False

        soup = BeautifulSoup(res.text, "html.parser")
        all_text = str(soup)
        
        # Find Links
        for link in soup.find_all("a", href=True):
            absolute = urljoin(url, link['href'])
            if urlparse(absolute).netloc == domain:
                links.add(normalize_url(absolute))
                
        # Find Integrations
        for name, patterns in integration_patterns.items():
            if any(re.search(p, all_text, re.I) for p in patterns):
                integrations_found.add(name)
                
        # Find Forms
        for form in soup.find_all("form"):
            fields = len(form.find_all(['input', 'textarea', 'select']))
            if fields > 0:
                forms_found.append({"Page URL": url, "Form Action": form.get('action', 'None'), "Fields Count": fields})
                
        # Find Calculators
        if any(kw in all_text.lower() for kw in ['calculate', 'calculator', 'estimate', 'mortgage', 'bmi']):
            has_calculator = True

        # MEMORY FIX: Instantly delete the heavy HTML objects to free up RAM for the next page
        del res, soup, all_text 

    except Exception:
        pass

    return url, links, integrations_found, forms_found, has_calculator

# -----------------------------
# 3. PARALLEL CRAWLER
# -----------------------------
def crawl_site(start_url, max_workers=15, max_pages=5000):
    domain = urlparse(start_url).netloc
    start_url = normalize_url(start_url)
    
    visited = set()
    queued = {start_url} 
    edges = []
    queue = [start_url]
    
    all_integrations = defaultdict(set)
    all_forms = []
    calc_pages = []
    
    integration_patterns = {
        'Google Analytics': [r'google-analytics\.com', r'googletagmanager\.com'],
        'Facebook Pixel': [r'facebook\.com/tr', r'connect\.facebook\.net'],
        'Hotjar': [r'hotjar\.com'], 'Intercom': [r'intercom\.io'],
        'HubSpot': [r'hubspot\.com'], 'Mailchimp': [r'mailchimp\.com'],
        'Zendesk': [r'zdassets\.com'], 'LiveChat': [r'livechatinc\.com'],
        'Stripe': [r'stripe\.com'], 'PayPal': [r'paypal\.com']
    }
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    while queue and len(visited) < max_pages:
        batch = queue[:max_workers]
        queue = queue[max_workers:]
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = [executor.submit(fetch_and_analyze, u, domain, integration_patterns) for u in batch]
            
            for future in as_completed(futures):
                url, links, integrations, forms, has_calc = future.result()
                
                if url in visited: continue
                visited.add(url)
                
                for name in integrations: all_integrations[name].add(url)
                all_forms.extend(forms)
                if has_calc: calc_pages.append(url)
                
                for link in links:
                    edges.append((url, link))
                    if link not in visited and link not in queued:
                        queued.add(link)
                        queue.append(link)
                        
        progress = min(len(visited) / max_pages, 1.0)
        progress_bar.progress(progress)
        status_text.text(f"🔄 Crawled Pages: {len(visited)} / {max_pages} | RAM Optimized")
        
    progress_bar.empty()
    status_text.empty()
    return visited, edges, all_integrations, all_forms, calc_pages

# -----------------------------
# 4. METRICS ENGINE
# -----------------------------
def calculate_metrics(start_url, pages, edges):
    df_edges = pd.DataFrame(edges, columns=["From", "To"]) if edges else pd.DataFrame(columns=["From", "To"])
    linked_pages = set(df_edges["To"]) if not df_edges.empty else set()
    orphan_pages = set(pages) - linked_pages
    if start_url in orphan_pages: orphan_pages.remove(start_url) 

    depth_map = {start_url: 0}
    bfs_queue = [start_url]
    adj = defaultdict(list)
    for frm, to in edges: adj[frm].append(to)
        
    while bfs_queue:
        current = bfs_queue.pop(0)
        for neighbor in adj[current]:
            if neighbor not in depth_map:
                depth_map[neighbor] = depth_map[current] + 1
                bfs_queue.append(neighbor)
                
    avg_depth = sum(depth_map.values()) / len(depth_map) if depth_map else 0

    metrics = {
        "Total Pages Crawled": len(pages),
        "Total Internal Links": len(edges),
        "Orphan Pages": len(orphan_pages),
        "% Orphan Pages": round((len(orphan_pages)/len(pages))*100, 2) if pages else 0,
        "Avg Navigation Depth": round(avg_depth, 2)
    }
    return metrics, orphan_pages, df_edges, depth_map

# -----------------------------
# 5. EXCEL REPORT (LOW MEMORY MODE)
# -----------------------------
def generate_excel(pages, edges, metrics, depth_map, integrations, forms, calcs):
    output = BytesIO()
    
    df_pages = pd.DataFrame(list(pages), columns=["Page URL"])
    df_pages['Navigation Depth'] = df_pages['Page URL'].map(depth_map).fillna('Unreachable')
    
    df_edges = pd.DataFrame(edges, columns=["From", "To"]) if edges else pd.DataFrame(columns=["From", "To"])
    df_metrics = pd.DataFrame(list(metrics.items()), columns=["Metric", "Value"])
    df_int = pd.DataFrame([(k, len(v)) for k, v in integrations.items()], columns=["Integration", "Pages Found On"])
    df_forms = pd.DataFrame(forms) if forms else pd.DataFrame(columns=["Page URL", "Form Action", "Fields Count"])
    df_calcs = pd.DataFrame(calcs, columns=["Calculator Pages"]) if calcs else pd.DataFrame(columns=["Calculator Pages"])

    # MEMORY FIX: constant_memory=True prevents Pandas from crashing RAM on massive files
    with pd.ExcelWriter(output, engine="xlsxwriter", options={'constant_memory': True}) as writer:
        df_metrics.to_excel(writer, sheet_name="Dashboard", index=False)
        df_pages.to_excel(writer, sheet_name="Pages & Depth", index=False)
        df_edges.to_excel(writer, sheet_name="Link Graph", index=False)
        df_int.to_excel(writer, sheet_name="Integrations", index=False)
        df_forms.to_excel(writer, sheet_name="Forms", index=False)
        df_calcs.to_excel(writer, sheet_name="Calculators", index=False)

    output.seek(0)
    return output

# -----------------------------
# 6. WORD REPORT
# -----------------------------
def generate_word(metrics, insights, integrations):
    doc = Document()
    doc.add_heading("Enterprise IA & Tech Audit Report", 0)
    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph("This report analyzes the website's Information Architecture and identifies third-party integrations, forms, and calculators.")
    doc.add_heading("Key Metrics", 1)
    for k, v in metrics.items(): doc.add_paragraph(f"{k}: {v}")
    doc.add_heading("Key Insights", 1)
    for i in insights: doc.add_paragraph(i)
    doc.add_heading("Third-Party Integrations Found", 1)
    if integrations:
        for name, urls in integrations.items(): doc.add_paragraph(f"{name}: Found on {len(urls)} pages.")
    else: doc.add_paragraph("No major third-party integrations detected.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------
# 7. USER INTERFACE
# -----------------------------
st.title("🌐 Enterprise IA & Tech Audit Tool")
st.markdown("Crawl massive websites (up to 10,000+ pages) to map URL depth, find orphan pages, and detect integrations/forms.")

col1, col2 = st.columns([3, 1])
with col1: url = st.text_input("Enter Website URL (include https://)", "https://www.mediclinic.ae")
with col2: 
    # FIX: Limit removed! You can now type up to 100,000
    max_pages = st.number_input("Max Pages", min_value=10, max_value=100000, value=500)

workers = st.slider("Crawl Speed (Threads)", 5, 25, 12)

if st.button("🚀 Start Enterprise Audit"):
    if not url: st.warning("Please enter a URL")
    else:
        with st.spinner("Crawling site... (This may take several minutes for large sites)"):
            pages, edges, integrations, forms, calcs = crawl_site(url, max_workers=workers, max_pages=max_pages)
            
        if not pages: st.error("Could not crawl the website. Check the URL and try again.")
        else:
            metrics, orphan_pages, df_edges, depth_map = calculate_metrics(url, pages, edges)
            
            insights = []
            if metrics["% Orphan Pages"] > 30: insights.append("⚠️ High orphan pages indicate weak internal linking structure.")
            if metrics["Avg Navigation Depth"] > 4: insights.append("⚠️ Deep navigation increases user effort and impacts UX.")
            if not insights: insights.append("✅ IA structure appears reasonably healthy.")

            tab1, tab2, tab3, tab4, tab5 = st.tabs(["📊 Dashboard", "🔌 Integrations", "📝 Forms", "🧮 Calculators", "⚠️ Orphans"])
            with tab1:
                st.subheader("Metrics"); st.json(metrics)
                st.subheader("Insights")
                for i in insights: st.write(i)
            with tab2:
                st.subheader(f"Integrations Found ({len(integrations)})")
                for name, urls in integrations.items():
                    with st.expander(f"🔹 {name} ({len(urls)} pages)"): st.write(list(urls)[:10])
            with tab3:
                st.subheader(f"Forms Found ({len(forms)})")
                if forms: st.dataframe(pd.DataFrame(forms), use_container_width=True)
                else: st.write("No forms found.")
            with tab4:
                st.subheader(f"Calculators Found ({len(calcs)})")
                if calcs: st.write(calcs[:50])
                else: st.write("No calculators found.")
            with tab5:
                st.subheader(f"Orphan Pages ({len(orphan_pages)})")
                st.write(list(orphan_pages)[:100])
                
            st.subheader("📥 Download Full Reports")
            excel = generate_excel(pages, edges, metrics, depth_map, integrations, forms, calcs)
            st.download_button("Download Excel Report", excel, "IA_Tech_Report.xlsx", type="primary")
            word = generate_word(metrics, insights, integrations)
            st.download_button("Download Word Report", word, "IA_Tech_Report.docx")